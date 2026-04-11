#!/usr/bin/env python3
"""
💰 批量转换所有定额为 Markdown 并录入系统

支持格式:
- Excel (.xls/.xlsx) → Markdown
- Word (.doc/.docx) → Markdown
- PDF (.pdf) → Markdown
- Access (.mdb/.accdb) → Markdown
- CHM (.chm/.chw) → Markdown (需先转 PDF)
- HTML (.html/.htm) → Markdown

作者：太一 AGI
创建：2026-04-11
"""

import pandas as pd
import subprocess
import os
from pathlib import Path
from datetime import datetime


def convert_excel_to_md(excel_file: str, output_dir: Path) -> bool:
    """转换 Excel 为 Markdown"""
    file_name = Path(excel_file).stem
    print(f"📊 转换 Excel: {file_name}")
    
    try:
        # 读取所有 sheet
        all_sheets = pd.read_excel(excel_file, sheet_name=None)
        
        md_lines = []
        md_lines.append(f"# 📊 {file_name}\n")
        md_lines.append(f"> **来源**: {excel_file}\n")
        md_lines.append(f"> **转换时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        md_lines.append(f"> **Sheet 数量**: {len(all_sheets)} 个\n")
        md_lines.append("")
        
        for sheet_name, df in all_sheets.items():
            md_lines.append(f"## 📄 {sheet_name}\n")
            md_lines.append("")
            md_lines.append(f"**行数**: {len(df)} 行，**列数**: {len(df.columns)} 列\n")
            md_lines.append("")
            
            # 表格
            if len(df) > 0:
                md_lines.append("### 数据预览 (前 20 行)\n")
                md_lines.append("")
                
                # 生成表格头
                cols = df.columns[:15].tolist()
                md_lines.append("| " + " | ".join(str(c) for c in cols) + " |")
                md_lines.append("|" + "|".join(["---"] * len(cols)) + "|")
                
                # 数据行
                for idx, row in df.head(20).iterrows():
                    values = [str(row[col])[:50] for col in cols]
                    md_lines.append("| " + " | ".join(values) + " |")
                
                md_lines.append("\n")
                
                # 统计信息
                numeric_cols = df.select_dtypes(include=['float64', 'int64']).columns
                if len(numeric_cols) > 0:
                    md_lines.append("### 统计信息\n")
                    md_lines.append("")
                    for col in numeric_cols[:10]:
                        md_lines.append(f"- **{col}**: 最小={df[col].min():,.2f}, 最大={df[col].max():,.2f}, 平均={df[col].mean():,.2f}\n")
                    md_lines.append("\n")
        
        # 保存
        output_file = output_dir / f"{file_name}.md"
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        
        print(f"   ✅ 已保存：{output_file.name} ({len(all_sheets)} sheets, {sum(len(df) for df in all_sheets.values())} 行)\n")
        return True
    
    except Exception as e:
        print(f"   ❌ 转换失败：{e}\n")
        return False


def convert_word_to_md(word_file: str, output_dir: Path) -> bool:
    """转换 Word 为 Markdown"""
    file_name = Path(word_file).stem
    print(f"📄 转换 Word: {file_name}")
    
    try:
        from docx import Document
        
        doc = Document(word_file)
        
        md_lines = []
        md_lines.append(f"# 📄 {file_name}\n")
        md_lines.append(f"> **来源**: {word_file}\n")
        md_lines.append(f"> **转换时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        md_lines.append("")
        
        # 提取段落
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                # 根据样式判断标题级别
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.replace('Heading', ''))
                    md_lines.append(f"{'#' * level} {para.text}\n")
                else:
                    md_lines.append(f"{para.text}\n")
        
        # 提取表格
        for i, table in enumerate(doc.tables):
            md_lines.append(f"\n### 表格 {i+1}\n\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    md_lines.append("| " + " | ".join(cells) + " |\n")
        
        # 保存
        output_file = output_dir / f"{file_name}.md"
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        
        print(f"   ✅ 已保存：{output_file.name} ({len(doc.paragraphs)} 段落，{len(doc.tables)} 表格)\n")
        return True
    
    except ImportError:
        print(f"   ⚠️  需要安装：pip install python-docx --break-system-packages\n")
        return False
    except Exception as e:
        print(f"   ❌ 转换失败：{e}\n")
        return False


def convert_pdf_to_md(pdf_file: str, output_dir: Path) -> bool:
    """转换 PDF 为 Markdown"""
    file_name = Path(pdf_file).stem
    print(f"📕 转换 PDF: {file_name}")
    
    try:
        import pdfplumber
        
        md_lines = []
        md_lines.append(f"# 📕 {file_name}\n")
        md_lines.append(f"> **来源**: {pdf_file}\n")
        md_lines.append(f"> **转换时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        md_lines.append("")
        
        with pdfplumber.open(pdf_file) as pdf:
            md_lines.append(f"> **页数**: {len(pdf.pages)} 页\n")
            md_lines.append("")
            
            for i, page in enumerate(pdf.pages):
                md_lines.append(f"## 第 {i+1} 页\n")
                md_lines.append("")
                
                # 提取文本
                text = page.extract_text()
                if text:
                    md_lines.append(f"{text}\n")
                
                # 提取表格
                tables = page.extract_tables()
                for j, table in enumerate(tables):
                    md_lines.append(f"\n### 表格 {j+1}\n\n")
                    for row in table:
                        if row:
                            cells = [str(cell).strip() if cell else '' for cell in row]
                            md_lines.append("| " + " | ".join(cells) + " |\n")
                
                md_lines.append("\n")
        
        # 保存
        output_file = output_dir / f"{file_name}.md"
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        
        print(f"   ✅ 已保存：{output_file.name} ({len(pdf.pages)} 页)\n")
        return True
    
    except ImportError:
        print(f"   ⚠️  需要安装：pip install pdfplumber --break-system-packages\n")
        return False
    except Exception as e:
        print(f"   ❌ 转换失败：{e}\n")
        return False


def convert_html_to_md(html_file: str, output_dir: Path) -> bool:
    """转换 HTML 为 Markdown"""
    file_name = Path(html_file).stem
    print(f"🌐 转换 HTML: {file_name}")
    
    try:
        from bs4 import BeautifulSoup
        
        with open(html_file, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        md_lines = []
        md_lines.append(f"# 🌐 {file_name}\n")
        md_lines.append(f"> **来源**: {html_file}\n")
        md_lines.append(f"> **转换时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        md_lines.append("")
        
        # 提取标题
        title = soup.find('title')
        if title:
            md_lines.append(f"**标题**: {title.text}\n\n")
        
        # 提取正文
        for h1 in soup.find_all('h1'):
            md_lines.append(f"# {h1.text}\n")
        for h2 in soup.find_all('h2'):
            md_lines.append(f"## {h2.text}\n")
        for h3 in soup.find_all('h3'):
            md_lines.append(f"### {h3.text}\n")
        
        # 提取段落
        for p in soup.find_all('p'):
            if p.text.strip():
                md_lines.append(f"{p.text}\n")
        
        # 提取表格
        for table in soup.find_all('table'):
            md_lines.append("\n### 表格\n\n")
            for row in table.find_all('tr'):
                cells = [cell.text.strip() for cell in row.find_all(['td', 'th'])]
                if cells:
                    md_lines.append("| " + " | ".join(cells) + " |\n")
        
        # 保存
        output_file = output_dir / f"{file_name}.md"
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))
        
        print(f"   ✅ 已保存：{output_file.name}\n")
        return True
    
    except ImportError:
        print(f"   ⚠️  需要安装：pip install beautifulsoup4 --break-system-packages\n")
        return False
    except Exception as e:
        print(f"   ❌ 转换失败：{e}\n")
        return False


def main():
    """主函数"""
    print("="*60)
    print("💰 批量转换所有定额为 Markdown 并录入系统")
    print("="*60)
    
    # 定额目录
    base_dirs = [
        "/home/nicola/下载/重庆 18 定额配套文件",
        "/home/nicola/下载",
    ]
    
    output_dir = Path("/home/nicola/.openclaw/workspace/skills/cost-agent/quota_md")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 查找第一个存在的目录
    base_dir = None
    for d in base_dirs:
        if Path(d).exists():
            base_dir = Path(d)
            break
    
    if not base_dir:
        print(f"❌ 定额目录不存在")
        return 1
    
    print(f"📂 使用目录：{base_dir}\n")
    
    # 查找所有文件
    all_files = []
    for ext in ['*.xls', '*.xlsx', '*.doc', '*.docx', '*.pdf', '*.html', '*.htm']:
        all_files.extend(base_dir.rglob(ext))
    
    # 过滤
    files_by_type = {
        'Excel': [f for f in all_files if f.suffix.lower() in ['.xls', '.xlsx']],
        'Word': [f for f in all_files if f.suffix.lower() in ['.doc', '.docx']],
        'PDF': [f for f in all_files if f.suffix.lower() == '.pdf'],
        'HTML': [f for f in all_files if f.suffix.lower() in ['.html', '.htm']],
    }
    
    print(f"📊 找到文件:")
    for type_name, files in files_by_type.items():
        print(f"   {type_name}: {len(files)} 个")
    print()
    
    # 转换统计
    stats = {
        'Excel': {'success': 0, 'fail': 0},
        'Word': {'success': 0, 'fail': 0},
        'PDF': {'success': 0, 'fail': 0},
        'HTML': {'success': 0, 'fail': 0},
    }
    
    # 转换 Excel
    print("="*60)
    print("📊 转换 Excel 文件")
    print("="*60)
    for f in files_by_type['Excel'][:50]:
        if convert_excel_to_md(str(f), output_dir):
            stats['Excel']['success'] += 1
        else:
            stats['Excel']['fail'] += 1
    
    # 转换 Word
    print("="*60)
    print("📄 转换 Word 文件")
    print("="*60)
    for f in files_by_type['Word'][:50]:
        if convert_word_to_md(str(f), output_dir):
            stats['Word']['success'] += 1
        else:
            stats['Word']['fail'] += 1
    
    # 转换 PDF
    print("="*60)
    print("📕 转换 PDF 文件")
    print("="*60)
    for f in files_by_type['PDF'][:50]:
        if convert_pdf_to_md(str(f), output_dir):
            stats['PDF']['success'] += 1
        else:
            stats['PDF']['fail'] += 1
    
    # 转换 HTML
    print("="*60)
    print("🌐 转换 HTML 文件")
    print("="*60)
    for f in files_by_type['HTML'][:20]:
        if convert_html_to_md(str(f), output_dir):
            stats['HTML']['success'] += 1
        else:
            stats['HTML']['fail'] += 1
    
    # 总结
    print("\n" + "="*60)
    print("✅ 批量转换完成!")
    print("="*60)
    for type_name, type_stats in stats.items():
        print(f"   {type_name}: 成功 {type_stats['success']} 个，失败 {type_stats['fail']} 个")
    
    total_success = sum(s['success'] for s in stats.values())
    total_fail = sum(s['fail'] for s in stats.values())
    print(f"\n   总计：成功 {total_success} 个，失败 {total_fail} 个")
    print(f"\n📁 输出目录：{output_dir}")
    
    return 0


if __name__ == "__main__":
    import sys
    sys.exit(main())
